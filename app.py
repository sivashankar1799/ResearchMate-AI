"""
ResearchMate AI - Main Flask Application
A production-ready research assistant for summarization, notes,
quizzes, chatbot Q&A, PPT generation, and exports.
"""

import os
import json
import uuid
import datetime
from flask import (
    Flask, render_template, request, jsonify, session,
    send_file, redirect, url_for
)
from werkzeug.utils import secure_filename

# Import utility modules (single-responsibility units)
from utils.pdf_reader import extract_pdf_text, validate_pdf
from utils.summarizer import (
    generate_short_summary, generate_detailed_summary,
    extract_key_findings, extract_conclusions
)
from utils.notes_generator import generate_study_notes, generate_revision_notes
from utils.quiz_generator import generate_mcqs, generate_true_false
from utils.chatbot import answer_question
from utils.ppt_generator import build_presentation

# Export utilities
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt, RGBColor

# ---------------------------------------------------------------------------
# App configuration
# ---------------------------------------------------------------------------
app = Flask(__name__)
app.secret_key = os.environ.get("RESEARCHMATE_SECRET", "researchmate-ai-secret-key-change-me")

# Base directory of the project for portable, non-hardcoded paths
BASE_DIR = os.path.abspath(os.path.dirname(__file__))

# Configure directories using app.config (no hardcoded paths)
app.config["UPLOAD_FOLDER"] = os.path.join(BASE_DIR, "uploads")
app.config["EXPORT_FOLDER"] = os.path.join(BASE_DIR, "exports")
app.config["MODEL_FOLDER"] = os.path.join(BASE_DIR, "models")
app.config["STATS_FILE"] = os.path.join(BASE_DIR, "models", "stats.json")
app.config["ALLOWED_EXTENSIONS"] = {"pdf"}
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50 MB upload limit

# Ensure required directories exist on startup
for folder in (app.config["UPLOAD_FOLDER"], app.config["EXPORT_FOLDER"], app.config["MODEL_FOLDER"]):
    os.makedirs(folder, exist_ok=True)


# ---------------------------------------------------------------------------
# Lightweight JSON-based statistics persistence
# ---------------------------------------------------------------------------
def _default_stats():
    """Return a fresh stats dictionary structure."""
    return {
        "papers_uploaded": 0,
        "total_pages": 0,
        "quizzes_generated": 0,
        "notes_generated": 0,
        "summaries_generated": 0,
        "ppts_generated": 0,
        "papers": [],          # list of {name, pages, words}
        "feature_usage": {     # for the pie chart
            "summaries": 0,
            "notes": 0,
            "quizzes": 0,
            "chats": 0,
            "ppts": 0
        }
    }


def load_stats():
    """Load stats from JSON file, creating defaults if missing/corrupt."""
    path = app.config["STATS_FILE"]
    if not os.path.exists(path):
        return _default_stats()
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
            # Merge with defaults to be resilient to schema changes
            base = _default_stats()
            base.update(data)
            return base
    except (json.JSONDecodeError, OSError):
        # Graceful fallback on a corrupt stats file
        return _default_stats()


def save_stats(stats):
    """Persist stats dictionary to JSON file."""
    try:
        with open(app.config["STATS_FILE"], "w", encoding="utf-8") as f:
            json.dump(stats, f, indent=2)
    except OSError:
        # Non-fatal: stats persistence failing should not crash a request
        pass


def bump_feature(feature_key, counter_key=None):
    """Increment a feature usage counter and optionally a top-level counter."""
    stats = load_stats()
    if feature_key in stats["feature_usage"]:
        stats["feature_usage"][feature_key] += 1
    if counter_key and counter_key in stats:
        stats[counter_key] += 1
    save_stats(stats)
    return stats


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def allowed_file(filename):
    """Validate file extension against the allowed set."""
    return "." in filename and \
        filename.rsplit(".", 1)[1].lower() in app.config["ALLOWED_EXTENSIONS"]


def get_document_text():
    """Retrieve the currently-loaded document text from the session."""
    return session.get("document_text", "")


def require_document():
    """Return (text, error_response) tuple; error_response is None when OK."""
    text = get_document_text()
    if not text:
        return None, (jsonify({"success": False,
                               "error": "No document loaded. Please upload a PDF first."}), 400)
    return text, None


# ---------------------------------------------------------------------------
# Page routes (render templates)
# ---------------------------------------------------------------------------
@app.route("/")
def index():
    """Dashboard / upload home page."""
    stats = load_stats()
    return render_template("index.html", stats=stats,
                           document_loaded=bool(get_document_text()),
                           document_name=session.get("document_name", ""))


@app.route("/summary")
def summary_page():
    return render_template("summary.html",
                           document_loaded=bool(get_document_text()),
                           document_name=session.get("document_name", ""))


@app.route("/notes")
def notes_page():
    return render_template("notes.html",
                           document_loaded=bool(get_document_text()),
                           document_name=session.get("document_name", ""))


@app.route("/quiz")
def quiz_page():
    return render_template("quiz.html",
                           document_loaded=bool(get_document_text()),
                           document_name=session.get("document_name", ""))


@app.route("/chatbot")
def chatbot_page():
    return render_template("chatbot.html",
                           document_loaded=bool(get_document_text()),
                           document_name=session.get("document_name", ""))


@app.route("/ppt")
def ppt_page():
    return render_template("ppt.html",
                           document_loaded=bool(get_document_text()),
                           document_name=session.get("document_name", ""))


# ---------------------------------------------------------------------------
# API: Upload & text extraction
# ---------------------------------------------------------------------------
@app.route("/upload", methods=["POST"])
def upload():
    """Handle PDF upload, validate, extract text, persist stats."""
    try:
        # Missing-file error handling
        if "file" not in request.files:
            return jsonify({"success": False, "error": "No file part in request."}), 400

        file = request.files["file"]
        if file.filename == "":
            return jsonify({"success": False, "error": "No file selected."}), 400

        # File-type validation
        if not allowed_file(file.filename):
            return jsonify({"success": False, "error": "Invalid file type. Only PDF is allowed."}), 400

        filename = secure_filename(file.filename)
        # Make filename unique to avoid collisions in uploads/
        unique_name = f"{uuid.uuid4().hex[:8]}_{filename}"
        save_path = os.path.join(app.config["UPLOAD_FOLDER"], unique_name)
        file.save(save_path)

        # Validate it is a readable PDF
        if not validate_pdf(save_path):
            os.remove(save_path)
            return jsonify({"success": False, "error": "The uploaded file is not a valid PDF."}), 400

        # Extract text + metadata via PyMuPDF
        text, page_count, word_count = extract_pdf_text(save_path)

        if not text.strip():
            return jsonify({"success": False,
                            "error": "Could not extract text. The PDF may be scanned/image-only."}), 422

        # Store in session for subsequent feature requests
        session["document_text"] = text
        session["document_name"] = filename
        session["document_path"] = save_path
        session["page_count"] = page_count
        session["word_count"] = word_count

        # Update persistent stats
        stats = load_stats()
        stats["papers_uploaded"] += 1
        stats["total_pages"] += page_count
        stats["papers"].append({"name": filename, "pages": page_count, "words": word_count})
        # Keep only the most recent 20 papers for chart readability
        stats["papers"] = stats["papers"][-20:]
        save_stats(stats)

        return jsonify({
            "success": True,
            "filename": filename,
            "page_count": page_count,
            "word_count": word_count,
            "message": "PDF uploaded and processed successfully."
        })
    except Exception as exc:  # Comprehensive error handling
        return jsonify({"success": False, "error": f"Upload failed: {str(exc)}"}), 500


# ---------------------------------------------------------------------------
# API: Summarization
# ---------------------------------------------------------------------------
@app.route("/summarize", methods=["POST"])
def summarize():
    """Generate short summary, detailed summary, key findings, conclusions."""
    try:
        text, err = require_document()
        if err:
            return err

        short = generate_short_summary(text)
        detailed = generate_detailed_summary(text)
        findings = extract_key_findings(text)
        conclusions = extract_conclusions(text)

        # Persist for export
        session["last_summary"] = {
            "short": short,
            "detailed": detailed,
            "findings": findings,
            "conclusions": conclusions
        }
        bump_feature("summaries", "summaries_generated")

        return jsonify({
            "success": True,
            "short_summary": short,
            "detailed_summary": detailed,
            "key_findings": findings,
            "conclusions": conclusions
        })
    except Exception as exc:
        return jsonify({"success": False, "error": f"Summarization failed: {str(exc)}"}), 500


# ---------------------------------------------------------------------------
# API: Notes
# ---------------------------------------------------------------------------
@app.route("/notes", methods=["POST"])
def notes():
    """Generate study notes and revision notes."""
    try:
        text, err = require_document()
        if err:
            return err

        study = generate_study_notes(text)
        revision = generate_revision_notes(text)

        session["last_notes"] = {"study": study, "revision": revision}
        bump_feature("notes", "notes_generated")

        return jsonify({"success": True, "study_notes": study, "revision_notes": revision})
    except Exception as exc:
        return jsonify({"success": False, "error": f"Notes generation failed: {str(exc)}"}), 500


# ---------------------------------------------------------------------------
# API: Quiz
# ---------------------------------------------------------------------------
@app.route("/quiz", methods=["POST"])
def quiz():
    """Generate 5 MCQs and 5 True/False questions."""
    try:
        text, err = require_document()
        if err:
            return err

        mcqs = generate_mcqs(text, count=5)
        tf = generate_true_false(text, count=5)

        session["last_quiz"] = {"mcqs": mcqs, "true_false": tf}
        bump_feature("quizzes", "quizzes_generated")

        return jsonify({"success": True, "mcqs": mcqs, "true_false": tf})
    except Exception as exc:
        return jsonify({"success": False, "error": f"Quiz generation failed: {str(exc)}"}), 500


# ---------------------------------------------------------------------------
# API: Chatbot
# ---------------------------------------------------------------------------
@app.route("/chat", methods=["POST"])
def chat():
    """Answer a question strictly from the uploaded document."""
    try:
        text, err = require_document()
        if err:
            return err

        data = request.get_json(silent=True) or {}
        question = (data.get("question") or "").strip()
        if not question:
            return jsonify({"success": False, "error": "Empty question."}), 400

        answer = answer_question(text, question)
        bump_feature("chats")

        return jsonify({
            "success": True,
            "answer": answer,
            "timestamp": datetime.datetime.now().strftime("%H:%M")
        })
    except Exception as exc:
        return jsonify({"success": False, "error": f"Chat failed: {str(exc)}"}), 500


# ---------------------------------------------------------------------------
# API: PPT generation
# ---------------------------------------------------------------------------
@app.route("/generate-ppt", methods=["POST"])
def generate_ppt():
    """Build a .pptx from the document and return it as a download."""
    try:
        text, err = require_document()
        if err:
            return err

        title = session.get("document_name", "Research Presentation").rsplit(".", 1)[0]
        out_name = f"{secure_filename(title)}_{uuid.uuid4().hex[:6]}.pptx"
        out_path = os.path.join(app.config["EXPORT_FOLDER"], out_name)

        build_presentation(text, title, out_path)
        bump_feature("ppts", "ppts_generated")

        return send_file(out_path, as_attachment=True, download_name=f"{title}.pptx")
    except Exception as exc:
        return jsonify({"success": False, "error": f"PPT generation failed: {str(exc)}"}), 500


# ---------------------------------------------------------------------------
# Export: Summary -> PDF (ReportLab)
# ---------------------------------------------------------------------------
@app.route("/export-summary", methods=["POST"])
def export_summary():
    """Export the last-generated summary to a PDF file."""
    try:
        summary = session.get("last_summary")
        if not summary:
            return jsonify({"success": False, "error": "No summary to export. Generate one first."}), 400

        out_name = f"summary_{uuid.uuid4().hex[:6]}.pdf"
        out_path = os.path.join(app.config["EXPORT_FOLDER"], out_name)

        doc = SimpleDocTemplate(out_path, pagesize=letter,
                                title="ResearchMate AI Summary")
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("TitleX", parent=styles["Title"],
                                     textColor="#1a237e")
        heading = ParagraphStyle("HeadingX", parent=styles["Heading2"],
                                 textColor="#283593", spaceBefore=14)
        body = styles["BodyText"]

        flow = [Paragraph("ResearchMate AI — Document Summary", title_style),
                Spacer(1, 0.2 * inch)]

        flow.append(Paragraph("Short Summary", heading))
        flow.append(Paragraph(summary.get("short", ""), body))

        flow.append(Paragraph("Detailed Summary", heading))
        flow.append(Paragraph(summary.get("detailed", ""), body))

        flow.append(Paragraph("Key Findings", heading))
        findings = summary.get("findings", [])
        if findings:
            flow.append(ListFlowable(
                [ListItem(Paragraph(f, body)) for f in findings], bulletType="bullet"))
        else:
            flow.append(Paragraph("No key findings extracted.", body))

        flow.append(Paragraph("Conclusions", heading))
        flow.append(Paragraph(summary.get("conclusions", ""), body))

        doc.build(flow)
        return send_file(out_path, as_attachment=True, download_name="ResearchMate_Summary.pdf")
    except Exception as exc:
        return jsonify({"success": False, "error": f"Summary export failed: {str(exc)}"}), 500


# ---------------------------------------------------------------------------
# Export: Notes -> DOCX (python-docx)
# ---------------------------------------------------------------------------
@app.route("/export-notes", methods=["POST"])
def export_notes():
    """Export the last-generated notes to a DOCX file."""
    try:
        notes_data = session.get("last_notes")
        if not notes_data:
            return jsonify({"success": False, "error": "No notes to export. Generate notes first."}), 400

        out_name = f"notes_{uuid.uuid4().hex[:6]}.docx"
        out_path = os.path.join(app.config["EXPORT_FOLDER"], out_name)

        document = Document()
        # Title
        title = document.add_heading("ResearchMate AI — Study Notes", level=0)

        document.add_heading("Study Notes", level=1)
        for group in notes_data.get("study", []):
            document.add_heading(group.get("topic", "Topic"), level=2)
            for point in group.get("points", []):
                p = document.add_paragraph(point, style="List Bullet")

        document.add_heading("Revision Notes", level=1)
        for point in notes_data.get("revision", []):
            document.add_paragraph(point, style="List Bullet")

        document.save(out_path)
        return send_file(out_path, as_attachment=True, download_name="ResearchMate_Notes.docx")
    except Exception as exc:
        return jsonify({"success": False, "error": f"Notes export failed: {str(exc)}"}), 500


# ---------------------------------------------------------------------------
# Export: Quiz results -> PDF (ReportLab)
# ---------------------------------------------------------------------------
@app.route("/export-quiz", methods=["POST"])
def export_quiz():
    """Export quiz results (score + correct answers) to PDF."""
    try:
        data = request.get_json(silent=True) or {}
        quiz_data = session.get("last_quiz")
        if not quiz_data:
            return jsonify({"success": False, "error": "No quiz to export. Generate a quiz first."}), 400

        score = data.get("score", "N/A")
        total = data.get("total", "N/A")

        out_name = f"quiz_{uuid.uuid4().hex[:6]}.pdf"
        out_path = os.path.join(app.config["EXPORT_FOLDER"], out_name)

        doc = SimpleDocTemplate(out_path, pagesize=letter, title="ResearchMate AI Quiz Results")
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("TitleX", parent=styles["Title"], textColor="#1a237e")
        heading = ParagraphStyle("HeadingX", parent=styles["Heading2"], textColor="#283593")
        body = styles["BodyText"]

        flow = [Paragraph("ResearchMate AI — Quiz Results", title_style),
                Spacer(1, 0.15 * inch),
                Paragraph(f"<b>Score:</b> {score} / {total}", body),
                Spacer(1, 0.2 * inch),
                Paragraph("Multiple Choice Questions", heading)]

        for i, q in enumerate(quiz_data.get("mcqs", []), 1):
            flow.append(Paragraph(f"<b>Q{i}. {q['question']}</b>", body))
            for opt in q["options"]:
                flow.append(Paragraph(f"&nbsp;&nbsp;• {opt}", body))
            flow.append(Paragraph(f"<i>Correct: {q['answer']}</i>", body))
            flow.append(Spacer(1, 0.1 * inch))

        flow.append(Paragraph("True / False Questions", heading))
        for i, q in enumerate(quiz_data.get("true_false", []), 1):
            flow.append(Paragraph(f"<b>Q{i}. {q['statement']}</b>", body))
            flow.append(Paragraph(f"<i>Correct: {q['answer']}</i>", body))
            flow.append(Spacer(1, 0.1 * inch))

        doc.build(flow)
        return send_file(out_path, as_attachment=True, download_name="ResearchMate_QuizResults.pdf")
    except Exception as exc:
        return jsonify({"success": False, "error": f"Quiz export failed: {str(exc)}"}), 500


# ---------------------------------------------------------------------------
# API: Stats (for dashboard charts auto-refresh)
# ---------------------------------------------------------------------------
@app.route("/stats", methods=["GET"])
def stats():
    """Return current dashboard statistics as JSON."""
    try:
        return jsonify({"success": True, "stats": load_stats()})
    except Exception as exc:
        return jsonify({"success": False, "error": str(exc)}), 500


# ---------------------------------------------------------------------------
# Error handlers
# ---------------------------------------------------------------------------
@app.errorhandler(413)
def too_large(_e):
    return jsonify({"success": False, "error": "File too large. Max 50MB."}), 413


@app.errorhandler(404)
def not_found(_e):
    return render_template("index.html", stats=load_stats(),
                           document_loaded=bool(get_document_text()),
                           document_name=session.get("document_name", "")), 404


if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)